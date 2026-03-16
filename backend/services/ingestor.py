import hashlib
import logging
import re
from pathlib import Path
from typing import Optional
import pdfplumber
from PyPDF2 import PdfReader
import docx

logger = logging.getLogger(__name__)


class IngestionError(Exception):
    pass


class ScannedPDFError(IngestionError):
    pass


class FileTooLargeError(IngestionError):
    pass


class InvalidFileError(IngestionError):
    pass


MAGIC_BYTES = {
    'pdf': b'%PDF',
    'docx': b'PK\x03\x04',
    'txt': None,
}

MAX_FILE_SIZE = 50 * 1024 * 1024


def validate_file(file_path: Path) -> None:
    if not file_path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
    
    if file_path.stat().st_size > MAX_FILE_SIZE:
        raise FileTooLargeError(f"Arquivo excede {MAX_FILE_SIZE // (1024*1024)}MB")
    
    ext = file_path.suffix.lower().lstrip('.')
    if ext not in MAGIC_BYTES:
        raise InvalidFileError(f"Extensão não suportada: {ext}")
    
    with open(file_path, 'rb') as f:
        header = f.read(4)
    
    if ext == 'pdf' and not header.startswith(MAGIC_BYTES['pdf']):
        raise InvalidFileError("Arquivo não é um PDF válido")
    elif ext == 'docx' and not header.startswith(MAGIC_BYTES['docx']):
        raise InvalidFileError("Arquivo não é um DOCX válido")


def extract_text_from_pdf(file_path: Path) -> str:
    text = ""
    page_texts = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                page_texts.append(page_text)
                text += page_text + "\n"
    except Exception as e:
        logger.warning(f"pdfplumber falhou: {e}, tentando PyPDF2")
    
    if not text.strip():
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                page_texts.append(page_text)
                text += page_text + "\n"
        except Exception as e:
            raise IngestionError(f"Erro ao extrair PDF: {e}")
    
    if not text.strip():
        raise ScannedPDFError(
            "PDF parece ser escaneado (sem texto extraível). "
            "Por favor, converta para PDF com texto selecionável."
        )
    
    text = _clean_extracted_text(text, page_texts)
    
    return text


def _clean_extracted_text(text: str, page_texts: list) -> str:
    lines = text.split('\n')
    
    if len(page_texts) > 1:
        first_lines = set()
        last_lines = set()
        
        for pt in page_texts[:3]:
            if pt:
                pl = pt.split('\n')
                if len(pl) >= 3:
                    first_lines.update(pl[:3])
        
        for pt in page_texts[-3:]:
            if pt:
                pl = pt.split('\n')
                if len(pl) >= 3:
                    last_lines.update(pl[-3:])
        
        cleaned_lines = []
        for line in lines:
            if line.strip() in first_lines or line.strip() in last_lines:
                if len(line.strip()) < 80:
                    continue
            cleaned_lines.append(line)
        text = '\n'.join(cleaned_lines)
    
    text = re.sub(r'\[\d+\]', '', text)
    text = re.sub(r'\([A-Za-z]+,\s*\d{4}\)', '', text)
    text = re.sub(r'Fig\.\s*\d+', '', text)
    text = re.sub(r'Tabela\s*\d+', '', text)
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
    
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def extract_text_from_docx(file_path: Path) -> str:
    try:
        doc = docx.Document(file_path)
        text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        raise IngestionError(f"Erro ao extrair DOCX: {e}")


def extract_text_from_txt(file_path: Path) -> str:
    encodings = ['utf-8', 'latin-1', 'cp1252']
    
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                text = f.read()
            break
        except UnicodeDecodeError:
            continue
    else:
        raise IngestionError("Não foi possível decodificar o arquivo")
    
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def extract_text(file_path: Path) -> str:
    validate_file(file_path)
    
    ext = file_path.suffix.lower().lstrip('.')
    
    if ext == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif ext == 'docx':
        text = extract_text_from_docx(file_path)
    elif ext == 'txt':
        text = extract_text_from_txt(file_path)
    else:
        raise InvalidFileError(f"Extensão não suportada: {ext}")
    
    if not text or len(text.strip()) < 50:
        raise IngestionError("Texto extraído muito curto ou vazio")
    
    logger.info(f"Texto extraído: {len(text)} caracteres de {file_path.name}")
    
    return text


def compute_text_hash(text: str) -> str:
    return hashlib.sha256(text.encode('utf-8')).hexdigest()


def ingest_file(file_path: Path) -> dict:
    text = extract_text(file_path)
    text_hash = compute_text_hash(text)
    
    return {
        'text': text,
        'char_count': len(text),
        'text_hash': text_hash,
        'file_type': file_path.suffix.lower().lstrip('.'),
    }
